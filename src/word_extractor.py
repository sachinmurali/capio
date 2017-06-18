import requests
import json
import time
from docx import Document
from docx.shared import RGBColor

class WordExtractor:
	def __init__(self):
		self.url = "https://api.capio.ai/v1/speech/transcript/"
		self.key = "262ac9a0c9ba4d179aad4c0b9b02120a"
		self.transcript_id = "593f237fbcae700012ba8fcd"

	def set_transcript_id(self, transcript_id):
		self.transcript_id = transcript_id

	def get_transcript(self):
		res = requests.get(self.url + self.transcript_id, headers={"apiKey" : self.key})
		return res

	def write_transcript(self):
		document = Document()
		transcripts = self.get_transcript().json()
		
		for ts in transcripts:		
			result = ts['result'][0]
			words = result['alternative'][0]['words']
			from_time = words[0]['from']
			start = self.get_start_time(from_time)

			s = document.add_paragraph()
			s.add_run(start + '\t').font.color.rgb = RGBColor(0x2d, 0x67, 0xc4)

			for word in words:
				if word['confidence'] < 0.75:
					s.add_run(word['word'] + ' ').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
				else:
					s.add_run(word['word'] + ' ')
		return document

	def get_start_time(self, from_time):
		minute, second = divmod(from_time, 60)
		hour, minute = divmod(minute, 60)
		return "%02d:%02d:%05.2f" % (hour,minute,second)

if __name__ == '__main__':
	wex = WordExtractor()
	doc = wex.write_transcript()
	doc.save(wex.transcript_id + '.docx')		